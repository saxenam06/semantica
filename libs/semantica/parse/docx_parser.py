"""
Word document parser for Semantica framework.

This module handles DOCX document parsing using python-docx
for text, formatting, and structure extraction.
"""

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..utils.exceptions import ProcessingError, ValidationError
from ..utils.logging import get_logger


@dataclass
class DocxSection:
    """Document section representation."""
    
    heading: str
    level: int
    content: str
    paragraphs: List[str] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)


@dataclass
class DocxMetadata:
    """Document metadata representation."""
    
    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    keywords: Optional[str] = None
    comments: Optional[str] = None
    category: Optional[str] = None
    created: Optional[str] = None
    modified: Optional[str] = None
    last_modified_by: Optional[str] = None


class DOCXParser:
    """DOCX document parser."""
    
    def __init__(self, **config):
        """
        Initialize DOCX parser.
        
        Args:
            **config: Parser configuration
        """
        self.logger = get_logger("docx_parser")
        self.config = config
    
    def parse(self, file_path: Union[str, Path], **options) -> Dict[str, Any]:
        """
        Parse DOCX document.
        
        Args:
            file_path: Path to DOCX file
            **options: Parsing options:
                - extract_formatting: Whether to extract formatting (default: False)
                - extract_tables: Whether to extract tables (default: True)
                - extract_comments: Whether to extract comments (default: False)
                
        Returns:
            dict: Parsed document data
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise ValidationError(f"DOCX file not found: {file_path}")
        
        if not file_path.suffix.lower() in ['.docx', '.doc']:
            raise ValidationError(f"File is not a DOCX: {file_path}")
        
        try:
            doc = Document(str(file_path))
            
            # Extract metadata
            metadata = self._extract_metadata(doc)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extract tables
            tables = []
            if options.get("extract_tables", True):
                for table in doc.tables:
                    table_data = self._extract_table(table)
                    tables.append(table_data)
            
            # Extract structure
            sections = self._extract_sections(doc, options)
            
            # Extract comments if requested
            comments = []
            if options.get("extract_comments", False):
                comments = self._extract_comments(doc)
            
            # Extract formatting if requested
            formatting_info = None
            if options.get("extract_formatting", False):
                formatting_info = self._extract_formatting(doc)
            
            return {
                "metadata": metadata.__dict__,
                "full_text": "\n\n".join(paragraphs),
                "paragraphs": paragraphs,
                "tables": tables,
                "sections": [s.__dict__ for s in sections],
                "comments": comments,
                "formatting": formatting_info,
                "total_paragraphs": len(paragraphs),
                "total_tables": len(tables)
            }
            
        except Exception as e:
            self.logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise ProcessingError(f"Failed to parse DOCX: {e}")
    
    def extract_text(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from DOCX.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            str: Extracted text
        """
        result = self.parse(file_path, extract_tables=False, extract_comments=False, extract_formatting=False)
        return result["full_text"]
    
    def extract_tables(self, file_path: Union[str, Path]) -> List[Dict[str, Any]]:
        """
        Extract tables from DOCX.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            list: Extracted tables
        """
        result = self.parse(file_path, extract_formatting=False, extract_comments=False)
        return result["tables"]
    
    def _extract_table(self, table: Table) -> Dict[str, Any]:
        """Extract data from table."""
        table_data = {
            "rows": [],
            "row_count": len(table.rows),
            "col_count": len(table.columns) if table.rows else 0
        }
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data["rows"].append(row_data)
        
        return table_data
    
    def _extract_sections(self, doc: DocxDocument, options: Dict[str, Any]) -> List[DocxSection]:
        """Extract document sections based on headings."""
        sections = []
        current_section = None
        
        for para in doc.paragraphs:
            # Check if paragraph is a heading
            if para.style.name.startswith('Heading'):
                # Get heading level
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                
                # Save previous section if exists
                if current_section:
                    sections.append(current_section)
                
                # Create new section
                current_section = DocxSection(
                    heading=para.text.strip(),
                    level=level,
                    content="",
                    paragraphs=[],
                    tables=[]
                )
            elif current_section:
                # Add paragraph to current section
                if para.text.strip():
                    current_section.paragraphs.append(para.text.strip())
                    current_section.content += para.text.strip() + "\n\n"
        
        # Add last section
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def _extract_metadata(self, doc: DocxDocument) -> DocxMetadata:
        """Extract document metadata."""
        core_props = doc.core_properties
        
        return DocxMetadata(
            title=core_props.title,
            author=core_props.author,
            subject=core_props.subject,
            keywords=core_props.keywords,
            comments=core_props.comments,
            category=core_props.category,
            created=str(core_props.created) if core_props.created else None,
            modified=str(core_props.modified) if core_props.modified else None,
            last_modified_by=core_props.last_modified_by
        )
    
    def _extract_comments(self, doc: DocxDocument) -> List[Dict[str, Any]]:
        """Extract comments from document."""
        comments = []
        # Note: python-docx doesn't directly support comments
        # This would require additional processing
        return comments
    
    def _extract_formatting(self, doc: DocxDocument) -> Dict[str, Any]:
        """Extract formatting information."""
        formatting = {
            "paragraphs": []
        }
        
        for para in doc.paragraphs:
            para_formatting = {
                "text": para.text,
                "style": para.style.name,
                "alignment": str(para.alignment) if para.alignment else None,
                "runs": []
            }
            
            for run in para.runs:
                run_formatting = {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_name": run.font.name if run.font.name else None,
                    "font_size": run.font.size.pt if run.font.size else None
                }
                para_formatting["runs"].append(run_formatting)
            
            formatting["paragraphs"].append(para_formatting)
        
        return formatting
